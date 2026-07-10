from uuid import UUID
from sqlalchemy.orm import Session
from fastapi import HTTPException, Depends
from core.supabase import get_db
from shared.database.model_resume import ResumeModel
import cloudinary.uploader
from sqlalchemy.orm.attributes import flag_modified
import urllib.request
import io
import docx
from .resume_ats_schema import KeywordAnalysisItem
from shared.database.model_ats_resume import AtsResumeModel
from core.langchain import get_llm
from pydantic import BaseModel
from langchain_core.prompts import PromptTemplate



class ResumeAtsService:
    def __init__(self, db: Session = Depends(get_db)):
        self.db = db

    def get_resume(self, user_id: UUID, job_id: str):
        ats_resume = self.db.query(AtsResumeModel).filter(
            AtsResumeModel.user_id == user_id,
            AtsResumeModel.job_id == job_id
        ).first()
        
        if not ats_resume:
            raise HTTPException(status_code=404, detail="ATS Resume not found for this job")
            
        return ats_resume

    def update_resume(self, user_id: UUID, job_id: str, old_text: str, new_text: str):
        ats_resume = self.get_resume(user_id, job_id)
        
        extracted_data = ats_resume.extracted_data or {}
        if not isinstance(extracted_data, dict):
            extracted_data = {}
            
        current_text = extracted_data.get("text", "")
        updated_text = current_text.replace(old_text, new_text)
        
        try:
            req = urllib.request.Request(ats_resume.cloudinary_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                original_bytes = response.read()

            doc = docx.Document(io.BytesIO(original_bytes))
            
            self._update_docx_content(doc, old_text, updated_text)
            
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            content = docx_stream.getvalue()
            
            # Extract original public_id from Cloudinary URL to overwrite it
            import re
            match = re.search(r'/upload/(?:v\d+/)?(.+)$', ats_resume.cloudinary_url)
            original_public_id = match.group(1) if match else f"user_{user_id}_updated_resume.docx"
            
            upload_result = cloudinary.uploader.upload(
                content, 
                resource_type="raw", 
                public_id=original_public_id,
                invalidate=True
            )
            new_cloudinary_url = upload_result.get("secure_url")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Cloudinary upload failed: {str(e)}")
            
        ats_resume.cloudinary_url = new_cloudinary_url
        
        extracted_data["text"] = updated_text
        if "filename" not in extracted_data:
            extracted_data["filename"] = "updated_resume.docx"
        
        ats_resume.extracted_data = extracted_data
        flag_modified(ats_resume, "extracted_data")
        
        self.db.commit()
        self.db.refresh(ats_resume)
        
        return ats_resume

    def _update_docx_content(self, doc, old_text: str, updated_text: str):
        old_lines = old_text.split('\n')
        new_lines = updated_text.split('\n')

        for i, p in enumerate(doc.paragraphs):
            if i < len(old_lines) and i < len(new_lines) and old_lines[i] != new_lines[i]:
                self._replace_paragraph_text(p, old_lines[i], new_lines[i])
                            
        # Add any new lines at the end
        if len(new_lines) > len(doc.paragraphs):
            for i in range(len(doc.paragraphs), len(new_lines)):
                doc.add_paragraph(new_lines[i])

    def _replace_paragraph_text(self, p, s_old: str, s_new: str):
        changed_old, changed_new = self._get_string_difference(s_old, s_new)
        
        replaced = False
        if changed_old:
            for run in p.runs:
                if changed_old in run.text:
                    run.text = run.text.replace(changed_old, changed_new)
                    replaced = True
                    
        # Fallback if the strict replacement failed or spanned multiple runs
        if not replaced:
            if p.runs:
                p.runs[0].text = s_new
                for j in range(1, len(p.runs)):
                    p.runs[j].text = ""
            else:
                p.text = s_new

    def _get_string_difference(self, s_old: str, s_new: str):
        prefix_len = 0
        while prefix_len < len(s_old) and prefix_len < len(s_new) and s_old[prefix_len] == s_new[prefix_len]:
            prefix_len += 1
            
        suffix_len = 0
        while suffix_len < len(s_old) - prefix_len and suffix_len < len(s_new) - prefix_len and s_old[-1 - suffix_len] == s_new[-1 - suffix_len]:
            suffix_len += 1
            
        changed_old = s_old[prefix_len : len(s_old) - suffix_len]
        changed_new = s_new[prefix_len : len(s_new) - suffix_len]
        return changed_old, changed_new

class ExtractedKeyword(BaseModel):
    keyword: str
    priority: str  # high, medium, low
    weightage: int

class RawKeywordsList(BaseModel):
    keywords: list[ExtractedKeyword]

class KeywordEditSuggestion(BaseModel):
    keyword: str
    old_sentence: str
    new_sentence: str

class KeywordEditSuggestionsList(BaseModel):
    suggestions: list[KeywordEditSuggestion]

class KeywordList(BaseModel):
    analysis_data: list[KeywordAnalysisItem]

class ResumeAtsService:
    def __init__(self, db: Session = Depends(get_db)):
        self.db = db

    def get_resume(self, user_id: UUID, job_id: str):
        ats_resume = self.db.query(AtsResumeModel).filter(
            AtsResumeModel.user_id == user_id,
            AtsResumeModel.job_id == job_id
        ).first()
        
        if not ats_resume:
            raise HTTPException(status_code=404, detail="ATS Resume not found for this job")
            
        return ats_resume

    def update_resume(self, user_id: UUID, job_id: str, old_text: str, new_text: str):
        ats_resume = self.get_resume(user_id, job_id)
        
        extracted_data = ats_resume.extracted_data or {}
        if not isinstance(extracted_data, dict):
            extracted_data = {}
            
        current_text = extracted_data.get("text", "")
        updated_text = current_text.replace(old_text, new_text)
        
        try:
            req = urllib.request.Request(ats_resume.cloudinary_url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                original_bytes = response.read()

            doc = docx.Document(io.BytesIO(original_bytes))
            
            self._update_docx_content(doc, old_text, updated_text)
            
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            content = docx_stream.getvalue()
            
            # Extract original public_id from Cloudinary URL to overwrite it
            import re
            match = re.search(r'/upload/(?:v\d+/)?(.+)$', ats_resume.cloudinary_url)
            original_public_id = match.group(1) if match else f"user_{user_id}_updated_resume.docx"
            
            upload_result = cloudinary.uploader.upload(
                content, 
                resource_type="raw", 
                public_id=original_public_id,
                invalidate=True
            )
            new_cloudinary_url = upload_result.get("secure_url")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Cloudinary upload failed: {str(e)}")
            
        ats_resume.cloudinary_url = new_cloudinary_url
        
        extracted_data["text"] = updated_text
        if "filename" not in extracted_data:
            extracted_data["filename"] = "updated_resume.docx"
        
        ats_resume.extracted_data = extracted_data
        flag_modified(ats_resume, "extracted_data")
        
        self.db.commit()
        self.db.refresh(ats_resume)
        
        return ats_resume

    def _update_docx_content(self, doc, old_text: str, updated_text: str):
        old_lines = old_text.split('\n')
        new_lines = updated_text.split('\n')

        for i, p in enumerate(doc.paragraphs):
            if i < len(old_lines) and i < len(new_lines) and old_lines[i] != new_lines[i]:
                self._replace_paragraph_text(p, old_lines[i], new_lines[i])
                            
        # Add any new lines at the end
        if len(new_lines) > len(doc.paragraphs):
            for i in range(len(doc.paragraphs), len(new_lines)):
                doc.add_paragraph(new_lines[i])

    def _replace_paragraph_text(self, p, s_old: str, s_new: str):
        changed_old, changed_new = self._get_string_difference(s_old, s_new)
        
        replaced = False
        if changed_old:
            for run in p.runs:
                if changed_old in run.text:
                    run.text = run.text.replace(changed_old, changed_new)
                    replaced = True
                    
        # Fallback if the strict replacement failed or spanned multiple runs
        if not replaced:
            if p.runs:
                p.runs[0].text = s_new
                for j in range(1, len(p.runs)):
                    p.runs[j].text = ""
            else:
                p.text = s_new

    def _get_string_difference(self, s_old: str, s_new: str):
        prefix_len = 0
        while prefix_len < len(s_old) and prefix_len < len(s_new) and s_old[prefix_len] == s_new[prefix_len]:
            prefix_len += 1
            
        suffix_len = 0
        while suffix_len < len(s_old) - prefix_len and suffix_len < len(s_new) - prefix_len and s_old[-1 - suffix_len] == s_new[-1 - suffix_len]:
            suffix_len += 1
            
        changed_old = s_old[prefix_len : len(s_old) - suffix_len]
        changed_new = s_new[prefix_len : len(s_new) - suffix_len]
        return changed_old, changed_new

    def analyze_job_keywords(self, user_id: UUID, job_description: str):
        latest_resume = self.db.query(ResumeModel).filter(
            ResumeModel.user_id == user_id
        ).order_by(ResumeModel.created_at.desc()).first()
        if not latest_resume:
            raise HTTPException(status_code=404, detail="Primary resume not found")
        resume_text = latest_resume.extracted_data.get("text", "") if isinstance(latest_resume.extracted_data, dict) else ""

        import os
        if not os.getenv("GROQ_API_KEY"):
            raise HTTPException(
                status_code=403,
                detail="GROQ_API_KEY is not set in backend configuration. Please add it to api/.env to run the ATS optimizer."
            )

        # Step 1: Extract all potential keywords from job description
        prompt_raw = PromptTemplate.from_template("""
        You are an expert ATS (Applicant Tracking System) optimizer.
        Analyze the following Job Description to extract the most important technical and non-technical keywords.
        Non-technical keywords should include soft skills, methodologies (like Agile/Scrum), and domain-specific roles or operations.
        
        Rules:
        1. Extract the most important technical and non-technical keywords from the Job Description.
        2. Assign a priority to each keyword: "high", "medium", or "low".
        3. Assign a weightage (integer) to each keyword indicating its importance. 
           CRITICAL: The sum of ALL weightages of the extracted keywords MUST EQUAL EXACTLY 100.
        
        Job Description:
        {job_description}
        """)

        try:
            raw_llm = get_llm().with_structured_output(RawKeywordsList)
            raw_chain = prompt_raw | raw_llm
            raw_keywords: RawKeywordsList = raw_chain.invoke({
                "job_description": job_description
            })
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract keywords from Job Description: {str(e)}"
            )

        # Step 2: Perform regex check to identify missing keywords
        import re
        matching_set = set()
        missing_keywords = []
        
        for kw in raw_keywords.keywords:
            escaped = re.escape(kw.keyword)
            try:
                is_match = bool(re.search(rf"(?<!\w){escaped}(?!\w)", resume_text, re.IGNORECASE))
            except Exception:
                is_match = kw.keyword.lower() in resume_text.lower()
                
            if is_match:
                matching_set.add(kw.keyword)
            else:
                missing_keywords.append(kw.keyword)

        # Step 3: Run LLM to get suggested edits for non-matching keywords
        suggestions_map = {}
        if missing_keywords:
            prompt_edit = PromptTemplate.from_template("""
            You are an expert resume editor and ATS optimizer.
            You are given a list of keywords that are missing from the user's Resume Text.
            For each missing keyword, find a relevant sentence from the Resume Text and rewrite it minimally to include the keyword.
            
            Rules:
            1. For each missing keyword, identify a sentence inside the Resume Text (old_sentence).
            2. The old_sentence MUST be an EXACT, case-sensitive substring from the provided Resume Text. Do not rewrite, summarize, or modify the old_sentence.
            3. Create a new_sentence by making the MINIMAL possible change to the old_sentence to include the keyword (e.g. inserting it in a list or appending it parenthetically). Keep the rest of the sentence identical.
            
            Missing Keywords:
            {missing_keywords}
            
            Resume Text:
            {resume_text}
            """)
            
            try:
                edit_llm = get_llm().with_structured_output(KeywordEditSuggestionsList)
                edit_chain = prompt_edit | edit_llm
                suggestions_result: KeywordEditSuggestionsList = edit_chain.invoke({
                    "missing_keywords": ", ".join(missing_keywords),
                    "resume_text": resume_text
                })
                
                if suggestions_result and suggestions_result.suggestions:
                    for s in suggestions_result.suggestions:
                        suggestions_map[s.keyword] = s
            except Exception as e:
                # Log the error but don't crash entirely; fallbacks will be used
                print(f"Failed to generate keyword edit suggestions: {str(e)}")

        # Step 4: Map keywords back to final KeywordAnalysisItem objects
        # Split resume text into sentences as a fallback option
        sentences = [s.strip() for s in re.split(r'[.!?\n]+', resume_text) if len(s.strip()) > 15]
        default_sentence = sentences[0] if sentences else "Professional experience."
        
        final_items = []
        for kw in raw_keywords.keywords:
            is_matching = kw.keyword in matching_set
            
            old_sentence = None
            new_sentence = None
            if not is_matching:
                sugg = suggestions_map.get(kw.keyword)
                if sugg:
                    # Validate that old_sentence is indeed a substring of resume_text
                    if sugg.old_sentence in resume_text:
                        old_sentence = sugg.old_sentence
                        new_sentence = sugg.new_sentence
                
                # Fallback if suggestion is missing or invalid
                if not old_sentence:
                    old_sentence = default_sentence
                    new_sentence = f"{old_sentence} (utilizing {kw.keyword})"
                    
            final_items.append(KeywordAnalysisItem(
                keyword=kw.keyword,
                priority=kw.priority,
                weightage=kw.weightage,
                is_matching=is_matching,
                old_sentence=old_sentence,
                new_sentence=new_sentence
            ))
            
        return KeywordList(analysis_data=final_items)

    def get_keyword_analysis(self, user_id: UUID, job_id: str):
        ats_resume = self.get_resume(user_id, job_id)
        return {
            "id": ats_resume.id,
            "analysis_data": ats_resume.analysis_data
        }

    def process_ats_resume(self, user_id: UUID, job_id: str, job_description: str):
        # 1. Check if it already exists
        existing_ats_resume = self.db.query(AtsResumeModel).filter(
            AtsResumeModel.user_id == user_id,
            AtsResumeModel.job_id == job_id
        ).first()

        if existing_ats_resume:
            return {"success": True, "message": "Resume already exists for this job, analysis skipped"}

        # 2. Fetch the latest resume
        latest_resume = self.db.query(ResumeModel).filter(
            ResumeModel.user_id == user_id
        ).order_by(ResumeModel.created_at.desc()).first()

        if not latest_resume:
            raise HTTPException(status_code=404, detail="Primary resume not found")

        # 3. Analyze job keywords (using the existing method, but we modified it to not save to DB)
        analysis_result = self.analyze_job_keywords(user_id, job_description)

        # 4. Save new AtsResumeModel
        new_ats_resume = AtsResumeModel(
            user_id=user_id,
            job_id=job_id,
            job_description=job_description,
            cloudinary_url=latest_resume.cloudinary_url,
            extracted_data=latest_resume.extracted_data,
            analysis_data=[item.model_dump() for item in analysis_result.analysis_data]
        )
        self.db.add(new_ats_resume)
        self.db.commit()

        return {"success": True, "message": "ATS resume processed successfully"}
